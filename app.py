import streamlit as st
import pandas as pd
from datetime import datetime, timedelta
import numpy as np
from scipy.optimize import newton
import calendar
from typing import Dict, Tuple, Optional
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Password protection - ADD THIS AT THE TOP
def check_password():
    """Returns True if password is correct"""
    def password_entered():
        if st.session_state["password"] == "AmicusLaw2025":
            st.session_state["password_correct"] = True
            del st.session_state["password"]  # Clear password from memory
        else:
            st.session_state["password_correct"] = False

    # Show password input if not authenticated
    if "password_correct" not in st.session_state:
        st.title("🏦 Guardian Ad Litem Calculator")
        st.write("**Access Required**")
        st.text_input("Enter Password:", type="password", 
                     on_change=password_entered, key="password")
        st.info("💡 Contact your administrator for the password")
        return False
    elif not st.session_state["password_correct"]:
        st.title("🏦 Guardian Ad Litem Calculator") 
        st.write("**Access Required**")
        st.text_input("Enter Password:", type="password", 
                     on_change=password_entered, key="password")
        st.error("❌ Incorrect password. Please try again.")
        return False
    else:
        return True

# Only show the app if password is correct
if check_password():

    # ==========================================
    # FINANCIAL CALCULATION FUNCTIONS
    # (DO NOT TOUCH THIS SECTION - WORKING FINANCIAL CODE)
    # ==========================================

    def xirr(cashflows, dates, guess=0.1):
        sorted_pairs = sorted(zip(dates, cashflows))
        dates = [pair[0] for pair in sorted_pairs]
        cashflows = [pair[1] for pair in sorted_pairs]
        
        first_date = dates[0]
        days = [(date - first_date).days for date in dates]
        
        def npv(rate):
            return sum(cf / ((1 + rate) ** (day / 365.0)) for cf, day in zip(cashflows, days))
        
        low = -0.99
        high = 10.0
        
        for _ in range(100):
            mid = (low + high) / 2.0
            npv_result = npv(mid)
            
            if abs(npv_result) < 1e-10:
                return mid
            elif npv_result > 0:
                low = mid
            else:
                high = mid
        
        return mid
    
    def find_treasury_bounds(duration_years: float) -> Tuple[float, float]:
        """
        Find the appropriate treasury bounds for interpolation.
        Available maturities: 0.25, 0.5, 1, 2, 3, 5, 7, 10, 20, 30 years
        For durations > 30 years, cap at 30-year rate (no extrapolation)
        """
        maturities = [0.25, 0.5, 1, 2, 3, 5, 7, 10, 20, 30]
        
        # Cap at 30-year rate for long durations
        if duration_years >= maturities[-1]:
            return maturities[-1], maturities[-1]  # 30, 30 (will result in flat 30Y rate)
        
        # Handle very short durations
        if duration_years <= maturities[0]:
            return maturities[0], maturities[1]  # 0.25, 0.5
        
        # Find bounds where duration sits between two maturities
        for i in range(len(maturities) - 1):
            if maturities[i] <= duration_years <= maturities[i + 1]:
                return maturities[i], maturities[i + 1]
        
        # Fallback (shouldn't happen)
        return 5, 7
    
    def get_treasury_series_info(maturity: float) -> Dict[str, str]:
        """
        Get FRED series information for a given maturity.
        Returns dict with series_id and display_name.
        """
        series_mapping = {
            0.25: {"series_id": "DGS3MO", "display_name": "3-Month"},
            0.5: {"series_id": "DGS6MO", "display_name": "6-Month"},
            1: {"series_id": "DGS1", "display_name": "1-Year"},
            2: {"series_id": "DGS2", "display_name": "2-Year"},
            3: {"series_id": "DGS3", "display_name": "3-Year"},
            5: {"series_id": "DGS5", "display_name": "5-Year"},
            7: {"series_id": "DGS7", "display_name": "7-Year"},
            10: {"series_id": "DGS10", "display_name": "10-Year"},
            20: {"series_id": "DGS20", "display_name": "20-Year"},
            30: {"series_id": "DGS30", "display_name": "30-Year"}
        }
        
        return series_mapping.get(maturity, {"series_id": "Unknown", "display_name": "Unknown"})
    
    def calculate_duration(payment_dates, payment_amounts, purchase_date, discount_rate):
        """
        Calculate weighted average duration of payments.
        Duration = Sum(PV × Years) / Sum(PV)
        """
        total_pv = 0
        total_time_weighted_pv = 0
        
        for payment_date, payment_amount in zip(payment_dates, payment_amounts):
            years = (payment_date - purchase_date).days / 365.0
            pv = payment_amount / ((1 + discount_rate) ** years)
            time_weighted_pv = pv * years
            
            total_pv += pv
            total_time_weighted_pv += time_weighted_pv
        
        if total_pv > 0:
            duration = total_time_weighted_pv / total_pv
            return duration
        else:
            return 0
    
    def calculate_excel_discount_rate(duration_years, lower_bound, upper_bound, lower_rate, upper_rate, spread):
        """
        Calculate discount rate using Excel's formula with user-provided treasury rates and spread
        Takes the actual bounds determined by find_treasury_bounds()
        """
        # Handle case where bounds are equal (duration >= 30 years)
        if upper_bound == lower_bound:
            # Use flat rate (no interpolation needed)
            discount_rate = lower_rate + spread
        else:
            # Normal interpolation between bounds
            discount_rate = ((duration_years - lower_bound) / (upper_bound - lower_bound) * (upper_rate - lower_rate)) + lower_rate + spread
        
        return discount_rate
    
    def calculate_wholesale_price(purchase_price, duration_years, total_payments, payment_dates, payment_amounts, purchase_date, excel_discount_rate):
        """
        Calculate wholesale price based on Excel formula in cell G5: C5+C13
        Uses the actual payment schedule for XNPV calculation (not simplified two-cash-flow model)
        Excel XNPV uses a 365-day year convention
        """
        if not payment_dates:
            return purchase_price
        
        # XNPV calculation using all actual payments
        # Cash flow 1: -purchase_price at time 0 (purchase date)
        # Cash flows 2+: individual payment amounts at their respective dates
        
        xnpv_value = -purchase_price  # Initial outflow
        
        # Add present value of each individual payment
        for payment_date, payment_amount in zip(payment_dates, payment_amounts):
            days_diff = (payment_date - purchase_date).days
            # Excel XNPV uses exact day count but 365-day year convention
            years_diff = days_diff / 365.0
            if years_diff >= 0:  # Only include future payments
                pv = payment_amount / ((1 + excel_discount_rate) ** years_diff)
                xnpv_value += pv
        
        # Wholesale price = Purchase price + XNPV
        wholesale_price = purchase_price + xnpv_value
        return wholesale_price
    
    def calculate_profit(wholesale_price, purchase_price, fixed_cost=6000):
        """
        Calculate profit based on Excel formula in cell G7: G5-C5-C15
        G7 = Wholesale Price - Purchase Price - Fixed Cost
        Fixed cost appears to be $6,000 based on Excel cell C15
        """
        profit = wholesale_price - purchase_price - fixed_cost
        return profit
    
    def calculate_competitor_quote(purchase_price, profit, target_profit=2500):
        """
        Calculate competitor quote based on Excel formula in cell C14: CEILING(C5+(G7-2500),50)
        This calculates what competitors might quote that would leave us with our target profit
        """
        import math
        competitor_quote = math.ceil((purchase_price + (profit - target_profit)) / 50) * 50
        return competitor_quote
    
    def generate_payment_schedule(num_payments, payment_amount, first_payment_date, last_payment_date, is_monthly):
        if num_payments == 1:
            return [first_payment_date], [payment_amount]
        
        dates = []
        
        if is_monthly:
            current_date = first_payment_date
            for i in range(num_payments):
                if i == 0:
                    dates.append(current_date)
                else:
                    if current_date.month == 12:
                        next_month = current_date.replace(year=current_date.year + 1, month=1)
                    else:
                        next_month = current_date.replace(month=current_date.month + 1)
                    try:
                        dates.append(next_month)
                        current_date = next_month
                    except ValueError:
                        last_day = calendar.monthrange(next_month.year, next_month.month)[1]
                        next_month = next_month.replace(day=last_day)
                        dates.append(next_month)
                        current_date = next_month
        else:
            current_date = first_payment_date
            for i in range(num_payments):
                if i == 0:
                    dates.append(current_date)
                else:
                    try:
                        next_year = current_date.replace(year=current_date.year + 1)
                        dates.append(next_year)
                        current_date = next_year
                    except ValueError:
                        next_year = current_date.replace(year=current_date.year + 1, day=28)
                        dates.append(next_year)
                        current_date = next_year
        
        amounts = [payment_amount] * num_payments
        return dates, amounts

    # ==========================================
    # WORD DOCUMENT GENERATION FUNCTIONS
    # (RESTORED FROM OLD CODE)
    # ==========================================

    def format_title_case_with_ordinals(text):
        """
        Format text to proper title case with special handling for ordinals and small words
        """
        if not text:
            return text
        
        # Define words that should remain lowercase (except when first word)
        small_words = {
            'a', 'an', 'and', 'as', 'at', 'but', 'by', 'for', 'in', 'nor', 'of', 
            'on', 'or', 'so', 'the', 'to', 'up', 'yet', 'with', 'from', 'into', 
            'onto', 'per', 'upon', 'via'
        }
        
        # Define ordinal patterns and their replacements
        ordinal_replacements = {
            '1st': '1ˢᵗ',
            '2nd': '2ⁿᵈ', 
            '3rd': '3ʳᵈ',
            '4th': '4ᵗʰ',
            '5th': '5ᵗʰ',
            '6th': '6ᵗʰ',
            '7th': '7ᵗʰ',
            '8th': '8ᵗʰ',
            '9th': '9ᵗʰ',
            '10th': '10ᵗʰ'
        }
        
        words = text.split()
        formatted_words = []
        
        for i, word in enumerate(words):
            # Remove punctuation for processing but remember it
            clean_word = word.strip('.,!?;:()[]{}"\'-')
            punctuation = word[len(clean_word):] if len(word) > len(clean_word) else ''
            
            # Check for ordinals first
            ordinal_found = False
            for ordinal, replacement in ordinal_replacements.items():
                if clean_word.lower() == ordinal.lower():
                    formatted_words.append(replacement + punctuation)
                    ordinal_found = True
                    break
            
            if not ordinal_found:
                # Apply title case rules
                if i == 0:  # First word is always capitalized
                    formatted_words.append(clean_word.capitalize() + punctuation)
                elif clean_word.lower() in small_words:
                    formatted_words.append(clean_word.lower() + punctuation)
                else:
                    formatted_words.append(clean_word.capitalize() + punctuation)
        
        return ' '.join(formatted_words)

    def create_tombs_maxwell_template(cause_number, factoring_company, courthouse, 
                                    report_title, sources_consulted, facts_section, 
                                    valuation_section):
        """
        Create Word document using the Tombs Maxwell template structure
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Book Antiqua'
        font.size = Pt(12)
        
        # Add cause number (bold, centered, uppercase, double-spaced)
        cause_para = doc.add_paragraph()
        cause_run = cause_para.add_run(f'CAUSE NO. {cause_number.upper()}')
        cause_run.bold = True
        cause_run.font.name = 'Book Antiqua'
        cause_run.font.size = Pt(12)
        cause_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cause_para.paragraph_format.line_spacing = 2.0  # Double spacing
        
        # Create IN RE table structure (2x1 table with thick center border)
        table = doc.add_table(rows=1, cols=2)
        
        # Remove all borders first, then add THINNER center border
        def set_table_borders(table):
            # Remove all borders first
            tbl = table._tbl
            for cell in table._cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is not None:
                    tcPr.remove(tcBorders)
            
            # Add THINNER center vertical border (sz="18" was at "24")
            left_cell = table.cell(0, 0)
            left_tc = left_cell._tc
            left_tcPr = left_tc.get_or_add_tcPr()
            left_tcBorders = parse_xml(r'<w:tcBorders %s><w:right w:val="single" w:sz="18" w:space="0" w:color="000000"/></w:tcBorders>' % nsdecls('w'))
            left_tcPr.append(left_tcBorders)
        
        set_table_borders(table)
        table.style = None
        
        # Left cell - IN RE: and factoring company (single-spaced with line breaks)
        in_re_cell = table.cell(0, 0)
        in_re_para = in_re_cell.paragraphs[0]
        in_re_para.clear()
        
        # Add "IN RE:" on first line
        in_re_run = in_re_para.add_run('IN RE:')
        in_re_run.bold = True
        in_re_run.font.name = 'Book Antiqua'
        in_re_run.font.size = Pt(12)
        
        # Add two line breaks and then factoring company
        in_re_para.add_run('\n\n\n')  # Skip 2 lines (3 \n total)
        
        company_run = in_re_para.add_run(factoring_company.upper())  # Convert to uppercase
        company_run.bold = True
        company_run.font.name = 'Book Antiqua'
        company_run.font.size = Pt(12)
        
        in_re_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        in_re_para.paragraph_format.line_spacing = 1.0  # Single spacing
        
        # Right cell - courthouse (double-spaced, LEFT JUSTIFIED with left margin indent)
        courthouse_cell = table.cell(0, 1)
        courthouse_para = courthouse_cell.paragraphs[0]
        courthouse_para.clear()
        
        # Add courthouse text
        courthouse_run = courthouse_para.add_run(courthouse.upper())  # Convert to uppercase
        courthouse_run.bold = True
        courthouse_run.font.name = 'Book Antiqua'
        courthouse_run.font.size = Pt(12)
        courthouse_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        courthouse_para.paragraph_format.line_spacing = 2.0  # Double spacing
        courthouse_para.paragraph_format.left_indent = Inches(0.35)  # Indent whole block 5 spaces
        
        # Add ONE space after table before REPORT OF GUARDIAN AD LITEM (remove space after)
        space_para = doc.add_paragraph()
        space_para.paragraph_format.line_spacing = 2.0
        space_para.paragraph_format.space_after = Pt(0)  # Remove space after paragraph
        
        # Add main heading with underline (double-spaced) - NO EXTRA SPACE BEFORE
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run('REPORT OF GUARDIAN AD LITEM')
        heading_run.bold = True  # Only this heading should be bold
        heading_run.underline = True
        heading_run.font.name = 'Book Antiqua'
        heading_run.font.size = Pt(12)
        heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_para.paragraph_format.line_spacing = 2.0
        
        # Add report title section (if provided)
        if report_title.strip():
            # Clean the text first - remove unwanted line breaks but preserve paragraph breaks
            cleaned_title = report_title.replace('\\\n', ' ').replace('\n\n', '|||PARAGRAPH_BREAK|||').replace('\n', ' ').replace('|||PARAGRAPH_BREAK|||', '\n\n')
            title_paragraphs = cleaned_title.split('\n\n')
            for title_para in title_paragraphs:
                if title_para.strip():
                    # Further clean each paragraph
                    clean_para = ' '.join(title_para.strip().split())
                    para = doc.add_paragraph(clean_para)
                    for run in para.runs:
                        run.font.name = 'Book Antiqua'
                        run.font.size = Pt(12)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.line_spacing = 2.0
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified
        
        # Add SOURCES CONSULTED section (double-spaced, NOT BOLD, INDENTED)
        sources_heading = doc.add_paragraph()
        sources_run = sources_heading.add_run('SOURCES CONSULTED:')
        sources_run.bold = False  # Remove bold
        sources_run.underline = True
        sources_run.font.name = 'Book Antiqua'
        sources_run.font.size = Pt(12)
        sources_heading.paragraph_format.line_spacing = 2.0
        sources_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sources_heading.paragraph_format.first_line_indent = Inches(0.5)  # Indent
        
        # Add sources paragraph (double-spaced, justified)
        if sources_consulted.strip():
            # Clean the text first - remove unwanted line breaks but preserve paragraph breaks
            cleaned_sources = sources_consulted.replace('\\\n', ' ').replace('\n\n', '|||PARAGRAPH_BREAK|||').replace('\n', ' ').replace('|||PARAGRAPH_BREAK|||', '\n\n')
            sources_paragraphs = cleaned_sources.split('\n\n')
            for source_para in sources_paragraphs:
                if source_para.strip():
                    # Further clean each paragraph
                    clean_para = ' '.join(source_para.strip().split())
                    para = doc.add_paragraph(clean_para)
                    for run in para.runs:
                        run.font.name = 'Book Antiqua'
                        run.font.size = Pt(12)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.line_spacing = 2.0
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified
        
        # Add FACTS section (double-spaced, NOT BOLD, INDENTED)
        facts_heading = doc.add_paragraph()
        facts_run = facts_heading.add_run('FACTS:')
        facts_run.bold = False  # Remove bold
        facts_run.underline = True
        facts_run.font.name = 'Book Antiqua'
        facts_run.font.size = Pt(12)
        facts_heading.paragraph_format.line_spacing = 2.0
        facts_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        facts_heading.paragraph_format.first_line_indent = Inches(0.5)  # Indent
        
        # Add facts paragraphs with indentation, double spacing, and justification
        if facts_section.strip():
            # Clean the text first - remove unwanted line breaks but preserve paragraph breaks
            cleaned_facts = facts_section.replace('\\\n', ' ').replace('\n\n', '|||PARAGRAPH_BREAK|||').replace('\n', ' ').replace('|||PARAGRAPH_BREAK|||', '\n\n')
            fact_paragraphs = cleaned_facts.split('\n\n')
            for fact_para in fact_paragraphs:
                if fact_para.strip():
                    # Further clean each paragraph
                    clean_para = ' '.join(fact_para.strip().split())
                    para = doc.add_paragraph(clean_para)
                    for run in para.runs:
                        run.font.name = 'Book Antiqua'
                        run.font.size = Pt(12)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.line_spacing = 2.0
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified
        
        # Add VALUATION and RECOMMENDATION section (double-spaced, NOT BOLD, INDENTED)
        valuation_heading = doc.add_paragraph()
        valuation_run = valuation_heading.add_run('VALUATION and RECOMMENDATION:')
        valuation_run.bold = False  # Remove bold
        valuation_run.underline = True
        valuation_run.font.name = 'Book Antiqua'
        valuation_run.font.size = Pt(12)
        valuation_heading.paragraph_format.line_spacing = 2.0
        valuation_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        valuation_heading.paragraph_format.first_line_indent = Inches(0.5)  # Indent
        
        # Add valuation paragraphs (double-spaced, justified)
        if valuation_section.strip():
            # Clean the text first - remove unwanted line breaks but preserve paragraph breaks
            cleaned_valuation = valuation_section.replace('\\\n', ' ').replace('\n\n', '|||PARAGRAPH_BREAK|||').replace('\n', ' ').replace('|||PARAGRAPH_BREAK|||', '\n\n')
            val_paragraphs = cleaned_valuation.split('\n\n')
            for val_para in val_paragraphs:
                if val_para.strip():
                    # Further clean each paragraph
                    clean_para = ' '.join(val_para.strip().split())
                    para = doc.add_paragraph(clean_para)
                    for run in para.runs:
                        run.font.name = 'Book Antiqua'
                        run.font.size = Pt(12)
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    para.paragraph_format.line_spacing = 2.0
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justified
        
        # Add several spaces before signature (double-spaced)
        space1 = doc.add_paragraph()
        space1.paragraph_format.line_spacing = 2.0
        
        # Create a container paragraph for right-aligned signature table
        container_para = doc.add_paragraph()
        container_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add signature block table (no borders, positioned on right side, WIDER)
        sig_table = doc.add_table(rows=3, cols=1)
        
        # Remove all borders from signature table
        def remove_all_table_borders(table):
            tbl = table._tbl
            for cell in table._cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is not None:
                    tcPr.remove(tcBorders)
        
        remove_all_table_borders(sig_table)
        sig_table.style = None
        
        # Position table to the right
        sig_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        # Set table width WIDER to accommodate email address
        sig_table.autofit = False
        for column in sig_table.columns:
            column.width = Inches(4.5)  # Increased from 3.5 to 4.5 inches for email
        
        # Move table to right by setting table properties
        tbl = sig_table._tbl
        tblPr = tbl.tblPr
        # Add table positioning with wider width
        tbl_pos = parse_xml(r'<w:tblW %s w:w="3240" w:type="dxa"/>' % nsdecls('w'))  # Increased width
        tblPr.append(tbl_pos)
        
        # Add table justification to right
        tbl_jc = parse_xml(r'<w:jc %s w:val="right"/>' % nsdecls('w'))
        tblPr.append(tbl_jc)
        
        # First row - "Respectfully submitted" and signature (single-spaced for signature block)
        first_row = sig_table.cell(0, 0)
        first_para = first_row.paragraphs[0]
        first_para.clear()
        
        # Add "Respectfully submitted," 
        resp_run = first_para.add_run("Respectfully submitted,")
        resp_run.font.name = 'Book Antiqua'
        resp_run.font.size = Pt(12)
        first_para.paragraph_format.line_spacing = 1.0
        first_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add THREE line breaks (one more single spaced line)
        first_para.add_run("\n\n\n")
        
        # Add signature with underline (NO space after - ensure no trailing space)
        sig_run = first_para.add_run("/s/ Joseph W. Tombs")
        sig_run.font.name = 'Book Antiqua'
        sig_run.font.size = Pt(12)
        sig_run.underline = True
        
        # Explicitly set space after paragraph to 0 to remove any trailing space
        first_para.paragraph_format.space_after = Pt(0)
        
        # Second row - empty for spacing
        empty_row = sig_table.cell(1, 0)
        empty_row.paragraphs[0].paragraph_format.line_spacing = 1.0
        
        # Third row - contact information (single-spaced for signature block)
        contact_cell = sig_table.cell(2, 0)
        contact_para = contact_cell.paragraphs[0]
        contact_para.clear()
        
        contact_text = """Joseph W. Tombs
TOMBS MAXWELL, LLP
State Bar No. 20116250
7021 Kewanee Ave. 7-102
Lubbock, TX 79424
Office (806) 698-1122
Tombs@tombsmaxwell.com"""
        
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.name = 'Book Antiqua'
        contact_run.font.size = Pt(12)
        contact_para.paragraph_format.line_spacing = 1.0
        contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Save to BytesIO object
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io

    # ==========================================
    # STREAMLIT APP INTERFACE
    # (MAIN APP STRUCTURE - SAFE TO MODIFY LAYOUT)
    # ==========================================
    
    # Streamlit App
    st.title("🏦 Amicus Law Offices, LLC")
    
    # Add tabs for Financial Analysis and Report Creation (reordered)
    tab1, tab2 = st.tabs(["💰 Financial Analysis", "📝 Report Creation"])
    
    with tab1:
        st.header("Financial Analysis")

        # NEW: Inform Nick of New Application Section - ADDED AT THE BEGINNING
        st.subheader("📧 Inform Nick of New Application")
        st.info("When you learn there is a new application, contact Nick to make the call and get the new facts for why the client is selling their payments.")
        
        # NEW: Lindsey Quote Section - ADDED AT THE BEGINNING
        st.subheader("📧 Get Quote from Lindsey")
        st.info("Before starting, email Lindsey at **lindsey@amicustrustcompany.com** to request a quote for this case.")
        
        # Create two columns for the quote inputs
        col1, col2 = st.columns(2)
        with col1:
            lindsey_quote = st.number_input(
                "Lindsey's Quote Amount ($):", 
                min_value=0.0, 
                value=0.0, 
                step=100.0, 
                format="%.2f",
                key="lindsey_quote_amount"
            )
        with col2:
            lindsey_irr = st.number_input(
                "Lindsey's IRR (%):", 
                min_value=0.0, 
                max_value=100.0, 
                value=0.0, 
                step=0.01, 
                format="%.2f",
                key="lindsey_irr_percent"
            )
        
        # Store Lindsey's data in session state
        st.session_state['lindsey_quote'] = lindsey_quote
        st.session_state['lindsey_irr'] = lindsey_irr / 100.0 if lindsey_irr > 0 else 0.0
        
        st.write("---")  # Separator line
        
        # Step 1: Number of payment groups
        st.subheader("Step 1: Select the number of payment groups")
        st.markdown("*Example 1: if the client is selling 5 payments of \\$7,000 and 2 payments of \\$4,000, you would select '2' because there are two uneven groups.*\n\n*Example 2: if the client is selling 165 payments of \\$500, you would select '1' because it is one large group of equal payments.*")
        num_groups = st.number_input("How many different payment groups are you selling?", min_value=1, value=1, step=1, key="financial_num_groups")

        # Collect data for each group
        all_payment_dates = []
        all_payment_amounts = []
        total_aggregate = 0

        for group_num in range(num_groups):
            if num_groups > 1:
                st.write("---")
                st.subheader(f"Group {group_num + 1} Details")
            
            step_offset = 1 if num_groups == 1 else 0
            
            st.write(f"**Step {2 + step_offset}: Payment Information - Group {group_num + 1}**")
            num_payments = st.number_input(f"How many payments in group {group_num + 1}?", min_value=1, value=1, step=1, key=f"financial_payments_{group_num}")

            # Purchase date selection - only show for the first group
            if group_num == 0:
                st.write(f"**Step {3 + step_offset}: Purchase Date**")
                use_today = st.radio("What date should be used for the purchase?", ["Today's date", "Different date"], key="financial_purchase_date_option")
                
                if use_today == "Today's date":
                    purchase_date = datetime.combine(datetime.now().date(), datetime.min.time())
                    st.write(f"**Purchase date: {purchase_date.strftime('%m/%d/%Y')}**")
                else:
                    custom_purchase_date = st.date_input(
                        "Select the purchase date:",
                        value=datetime.now().date(),
                        min_value=datetime.now().date() - timedelta(days=365*10),
                        max_value=datetime.now().date() + timedelta(days=365*10),
                        key="financial_custom_purchase_date"
                    )
                    purchase_date = datetime.combine(custom_purchase_date, datetime.min.time())
                    st.write(f"**Purchase date: {purchase_date.strftime('%m/%d/%Y')}**")
                
                # Update step numbers for subsequent steps
                step_offset += 1

            if num_payments > 1:
                st.write(f"**Step {3 + step_offset}: Payment Frequency - Group {group_num + 1}**")
                payment_frequency = st.radio(f"Are these annual or monthly payments?", ["Monthly", "Annual"], key=f"financial_frequency_{group_num}")
                is_monthly = payment_frequency == "Monthly"
            else:
                is_monthly = False

            st.write(f"**Step {4 + step_offset}: Payment Amount - Group {group_num + 1}**")
            payment_amount = st.number_input(f"How much is each payment in group {group_num + 1}?", min_value=0.01, value=10000.00, step=100.00, format="%.2f", key=f"financial_amount_{group_num}")

            group_aggregate = num_payments * payment_amount
            total_aggregate += group_aggregate
            
            st.write(f"**Group {group_num + 1} aggregate: \\${group_aggregate:,.2f}** ({num_payments}  payments × \\${payment_amount:,.2f} each)")

            st.write(f"**Step {5 + step_offset}: Payment Dates - Group {group_num + 1}**")
            first_payment_date = st.date_input(f"When will the first payment happen in group {group_num + 1}?", value=datetime.now().date() + timedelta(days=30), min_value=datetime.now().date(), max_value=datetime.now().date() + timedelta(days=365*50), key=f"financial_first_date_{group_num}")

            if num_payments > 1:
                last_payment_date = st.date_input(f"When will the last payment happen in group {group_num + 1}?", value=datetime.now().date() + timedelta(days=365), min_value=datetime.now().date(), max_value=datetime.now().date() + timedelta(days=365*50), key=f"financial_last_date_{group_num}")
                if last_payment_date <= first_payment_date:
                    st.error(f"Last payment date must be after first payment date in group {group_num + 1}!")
                    st.stop()
            else:
                last_payment_date = first_payment_date

            group_dates, group_amounts = generate_payment_schedule(num_payments, payment_amount, datetime.combine(first_payment_date, datetime.min.time()), datetime.combine(last_payment_date, datetime.min.time()), is_monthly)
            
            all_payment_dates.extend(group_dates)
            all_payment_amounts.extend(group_amounts)

        # Overall verification step
        st.write("---")
        st.subheader("⚠️ Overall Verification Step")
        st.write(f"**The total aggregate of ALL payments is ${total_aggregate:,.2f}**")
        if num_groups > 1:
            st.write("**Breakdown by group:**")
            for group_num in range(num_groups):
                num_payments_group = st.session_state.get(f"financial_payments_{group_num}", 1)
                amount_group = st.session_state.get(f"financial_amount_{group_num}", 10000.0)
                group_total = num_payments_group * amount_group
                st.write(f"• Group {group_num + 1}: {num_payments_group} payments × \\${amount_group:,.2f} = \\${group_total:,.2f}")

        aggregate_correct = st.radio("Is this total aggregate amount correct?", ["Select an option", "Yes, this is correct", "No, I need to update my numbers"], key="financial_aggregate_check")

        if aggregate_correct == "No, I need to update my numbers":
            st.warning("Please update your numbers above and check again.")
            st.stop()
        elif aggregate_correct == "Select an option":
            st.info("Please confirm if the total aggregate amount is correct before continuing.")
            st.stop()
        elif aggregate_correct == "Yes, this is correct":
            st.success("Great! Let's continue with the purchase price.")

        # Purchase price
        final_step = 7 if num_groups == 1 else 3 + num_groups * 4
        st.subheader(f"Step {final_step}: Purchase Price")
        purchase_price = st.number_input("How much is the factoring company buying ALL the payments for?", min_value=0.01, value=float(total_aggregate * 0.85), step=100.00, format="%.2f", key="financial_purchase_price")

        # Competitor analysis settings
        st.subheader(f"Step {final_step + 1}: Competitor Analysis")
        st.write("For competitor quote calculation, we need to set a target profit to determine competitive pricing.")
        use_default_target_profit = st.radio(
            "What target profit should we use for competitor quote calculation?", 
            ["Use $2,500 (default)", "Specify a different target profit"],
            key="financial_target_profit_choice"
        )

        if use_default_target_profit == "Use $2,500 (default)":
            target_profit = 2500
            st.write("**Using target profit: $2,500**")
        else:
            target_profit = st.number_input(
                "Enter the target profit amount:", 
                min_value=0.0, 
                value=2500.0, 
                step=100.0, 
                format="%.2f",
                key="financial_custom_target_profit"
            )
            st.write(f"**Using target profit: ${target_profit:,.2f}**")

        st.subheader("📊 Results & Analysis")

        # Sort all payments by date
        sorted_payment_pairs = sorted(zip(all_payment_dates, all_payment_amounts))
        payment_dates = [pair[0] for pair in sorted_payment_pairs]
        payment_amounts = [pair[1] for pair in sorted_payment_pairs]

        # Use the selected purchase date (either today or custom)
        cashflows = [-purchase_price] + payment_amounts
        dates = [purchase_date] + payment_dates

        irr_rate = xirr(cashflows, dates)

        if irr_rate is not None:
            # Calculate duration
            duration_years = calculate_duration(payment_dates, payment_amounts, purchase_date, irr_rate)
            
            # Determine which treasury bounds we need for the duration
            lower_bound, upper_bound = find_treasury_bounds(duration_years)
            
            # Display duration and treasury requirements
            st.write("---")
            st.subheader("🏛️ Treasury Rate Input Required")
            st.subheader(f"Duration: {duration_years:.2f} years")
            st.write(f"**Purchase date used: {purchase_date.strftime('%m/%d/%Y')}**")
            
            # Get series information for the bounds
            lower_series_info = get_treasury_series_info(lower_bound)
            upper_series_info = get_treasury_series_info(upper_bound)
            
            if lower_bound == upper_bound:
                st.write(f"**Need: {lower_series_info['display_name']} treasury rate** (duration ≥ 30 years, capped)")
                st.write(f"📄 **Get the current rate from:** https://fred.stlouisfed.org/series/{lower_series_info['series_id']}")
                
                # Single rate input
                manual_rate = st.number_input(
                    f"{lower_series_info['display_name']} Treasury Rate (%)",
                    min_value=0.0,
                    max_value=20.0,
                    value=4.0,
                    step=0.01,
                    format="%.2f",
                    help=f"Enter the most recent rate from the FRED page above",
                    key="financial_single_treasury_rate"
                )
                lower_rate = upper_rate = manual_rate / 100.0
                
            else:
                st.write(f"**Need: {lower_series_info['display_name']} and {upper_series_info['display_name']} treasury rates** for interpolation")
                st.write(f"📄 **Get the current rates from:**")
                st.write(f"• **{lower_series_info['display_name']}:** https://fred.stlouisfed.org/series/{lower_series_info['series_id']}")
                st.write(f"• **{upper_series_info['display_name']}:** https://fred.stlouisfed.org/series/{upper_series_info['series_id']}")
                
                # Two rate inputs
                col1, col2 = st.columns(2)
                with col1:
                    manual_lower = st.number_input(
                        f"{lower_series_info['display_name']} Treasury Rate (%)",
                        min_value=0.0,
                        max_value=20.0,
                        value=4.0,
                        step=0.01,
                        format="%.2f",
                        help=f"Enter the most recent rate from the FRED page above",
                        key="financial_lower_treasury_rate"
                    )
                with col2:
                    manual_upper = st.number_input(
                        f"{upper_series_info['display_name']} Treasury Rate (%)",
                        min_value=0.0,
                        max_value=20.0,
                        value=4.2,
                        step=0.01,
                        format="%.2f",
                        help=f"Enter the most recent rate from the FRED page above",
                        key="financial_upper_treasury_rate"
                    )
                
                lower_rate = manual_lower / 100.0
                upper_rate = manual_upper / 100.0
            
            # Spread input
            st.write("**📊 Spread Configuration**")
            use_default_spread = st.radio(
                "Would you like to use the default spread of 3.0%?", 
                ["Yes, use 3.0%", "No, I want to specify a different spread"],
                key="financial_spread_choice"
            )
            
            if use_default_spread == "Yes, use 3.0%":
                spread = 0.03
                st.write("**Using default spread: 3.0%**")
            else:
                spread_percentage = st.number_input(
                    "Enter the spread percentage:", 
                    min_value=0.0, 
                    max_value=10.0,
                    value=3.0, 
                    step=0.1, 
                    format="%.1f",
                    key="financial_custom_spread"
                )
                spread = spread_percentage / 100.0
                st.write(f"**Using custom spread: {spread_percentage:.1f}%**")
            
            # Calculate Excel discount rate using treasury rates and spread
            excel_discount_rate = calculate_excel_discount_rate(duration_years, lower_bound, upper_bound, lower_rate, upper_rate, spread)
            
            # Calculate wholesale price, profit, and competitor analysis
            total_payments = sum(payment_amounts)
            wholesale_price = calculate_wholesale_price(purchase_price, duration_years, total_payments, payment_dates, payment_amounts, purchase_date, excel_discount_rate)
            profit = calculate_profit(wholesale_price, purchase_price)
            competitor_quote = calculate_competitor_quote(purchase_price, profit, target_profit)
            
            # Calculate the profit if we match competitor's quote
            competitor_profit = calculate_profit(wholesale_price, competitor_quote)
            
            # Calculate XIRR for competitive scenario
            competitive_cashflows = [-competitor_quote] + payment_amounts
            competitive_dates = [purchase_date] + payment_dates
            competitive_irr = xirr(competitive_cashflows, competitive_dates)
            
            # Financial summary - Updated format
            st.write("**📈 Profit Analysis**")
            
            # Side-by-side profit calculations
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**💰 Factoring Company**")
                st.code(f"""
    Wholesale Price:       ${wholesale_price:,.2f}
    Less Purchase Price:  -${purchase_price:,.2f}
    Less Legal Costs:     -$6,000.00
                          ________________
    Profit:                ${profit:,.2f}
                """)
                st.markdown(f"""
                <div style="text-align: right; padding: 10px; border: 1px solid #ccc; border-radius: 5px; background-color: #f0f2f6;">
                    <div style="font-size: 14px; color: #666;">Factoring Company Discount Rate</div>
                    <div style="font-size: 24px; font-weight: bold; color: #333;">{irr_rate:.2%}</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                st.write("**🏢 Competitive Analysis**")
                competitive_irr_display = f"{competitive_irr:.2%}" if competitive_irr is not None else "N/A"
                st.code(f"""
    Wholesale Price:         ${wholesale_price:,.2f}
    Less Competitive Quote: -${competitor_quote:,.2f}
    Less Legal Costs:       -$6,000.00
                          ________________
    Profit:                ${competitor_profit:,.2f}
                """)
                st.markdown(f"""
                <div style="text-align: right; padding: 10px; border: 1px solid #ccc; border-radius: 5px; background-color: #f0f2f6;">
                    <div style="font-size: 14px; color: #666;">Competitive Quote Discount Rate</div>
                    <div style="font-size: 24px; font-weight: bold; color: #333;">{competitive_irr_display}</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Payment schedule
            st.write("**📅 Payment Schedule**")
            df = pd.DataFrame({
                'Payment Date': [d.strftime('%m/%d/%Y') for d in payment_dates], 
                'Payment Amount': [f"${amount:,.2f}" for amount in payment_amounts]
            })
            st.dataframe(df, hide_index=True)
            
            # Store financial data in session state for report creation
            st.session_state['financial_complete'] = True
            st.session_state['num_groups'] = num_groups
            st.session_state['total_aggregate'] = total_aggregate
            st.session_state['purchase_price'] = purchase_price
            st.session_state['all_payment_dates'] = all_payment_dates
            st.session_state['all_payment_amounts'] = all_payment_amounts
            
            # Detailed calculations (expandable) - MODIFIED TO INCLUDE LINDSEY'S DATA
            with st.expander("🔬 Detailed Calculations"):
                st.write("**Duration Calculation Details:**")
                duration_details = []
                total_pv = 0
                total_time_weighted_pv = 0
                
                for i, (payment_date, payment_amount) in enumerate(zip(payment_dates, payment_amounts)):
                    years = (payment_date - purchase_date).days / 365.0
                    pv = payment_amount / ((1 + irr_rate) ** years)
                    time_weighted_pv = pv * years
                    
                    total_pv += pv
                    total_time_weighted_pv += time_weighted_pv
                    
                    duration_details.append({
                        'Payment #': i + 1,
                        'Date': payment_date.strftime('%m/%d/%Y'),
                        'Years': f"{years:.3f}",
                        'Payment Amount': f"${payment_amount:,.2f}",
                        'Present Value': f"${pv:,.2f}",
                        'PV × Years': f"${time_weighted_pv:,.2f}"
                    })
                
                duration_df = pd.DataFrame(duration_details)
                st.dataframe(duration_df, hide_index=True)
                
                st.write(f"**Duration = ${total_time_weighted_pv:,.2f} ÷ ${total_pv:,.2f} = {duration_years:.6f} years**")
                
                # Calculate the XNPV components for display using Excel's discount rate and actual payments
                xnpv_initial = -purchase_price
                xnpv_payments = 0
                
                for payment_date, payment_amount in zip(payment_dates, payment_amounts):
                    days_diff = (payment_date - purchase_date).days
                    years_diff = days_diff / 365.0
                    if years_diff >= 0:
                        pv = payment_amount / ((1 + excel_discount_rate) ** years_diff)
                        xnpv_payments += pv
                
                xnpv_value = xnpv_initial + xnpv_payments
                
                # MODIFIED: Financial Calculations section now includes Lindsey's data
                st.write("**Financial Calculations:**")
                lindsey_quote_display = f"${st.session_state.get('lindsey_quote', 0):,.2f}" if st.session_state.get('lindsey_quote', 0) > 0 else "Not provided"
                lindsey_irr_display = f"{st.session_state.get('lindsey_irr', 0):.2%}" if st.session_state.get('lindsey_irr', 0) > 0 else "Not provided"
                
                st.code(f"""
    === LINDSEY'S QUOTE ===
    Lindsey's Quote Amount: {lindsey_quote_display}
    Lindsey's IRR: {lindsey_irr_display}
    
    === CALCULATED VALUES ===
    Total Payments: ${total_payments:,.2f}
    Purchase Price: ${purchase_price:,.2f}
    Duration: {duration_years:.3f} years
    Number of Payments: {len(payment_dates)}
    
    Treasury Rates Used:
      Lower Bound ({lower_bound}Y): {lower_rate:.4f} ({lower_rate:.2%})
      Upper Bound ({upper_bound}Y): {upper_rate:.4f} ({upper_rate:.2%})
    
    Excel Discount Rate: {excel_discount_rate:.4f} ({excel_discount_rate:.2%})
    (Formula: ((Duration-{lower_bound})/({upper_bound}-{lower_bound})*({upper_rate:.4f}-{lower_rate:.4f}))+{lower_rate:.4f}+{spread:.3f})
    
    Spread Used: {spread:.1%}
    
    XNPV Calculation (using actual payment schedule):
      PV of initial outflow: ${xnpv_initial:,.2f}
      PV of all payments: ${xnpv_payments:,.2f}
      XNPV Total: ${xnpv_value:,.2f}
    
    Wholesale Price: ${wholesale_price:,.2f} (Purchase Price + XNPV)
    Competitor Quote: ${competitor_quote:,.2f}
    Target Profit Used: ${target_profit:,.2f}
                """)

            # Navigation guidance
            st.write("---")
            st.write("### ✅ Financial Analysis Complete!")
            st.write("Ready to create your Guardian Ad Litem report? Click the **📝 Report Creation** tab above to continue.")

        else:
            st.error("Could not calculate XIRR. Please check your inputs.")

    # UPDATED TAB 2 - Report Creation (combines new simplified input with old perfect formatting)
    with tab2:
        st.header("Report Formatting Tool")
        st.write("Use this tool to format your AI-generated Guardian Ad Litem report into a properly formatted Word document.")
        
        # Check if financial analysis is complete
        if not st.session_state.get('financial_complete', False):
            st.warning("⚠️ Please complete the **💰 Financial Analysis** tab first.")
            st.info("👈 Click the **Financial Analysis** tab to get started!")
        
        st.write("---")
        
        # STEP 1: Basic case information (reduced to 3 questions)
        st.subheader("Step 1: Basic Case Information")
        
        cause_number = st.text_input(
            "What is the cause number?", 
            value="", 
            key="report_cause_number",
            help="Example: CC-2024-CV-0656"
        )
        
        factoring_company = st.text_input(
            "What is the factoring company's name on the application?", 
            value="", 
            key="report_factoring_company",
            help="Example: J.G. WENTWORTH ORIGINATIONS, LLC"
        )
        
        courthouse = st.text_input(
            "What is the court, number, county, state on the application?", 
            value="", 
            key="report_courthouse",
            help="Example: IN THE COUNTY COURT AT LAW NUMBER THREE (3) OF LUBBOCK COUNTY, TEXAS"
        )
        
        st.write("---")
        
        # STEP 2: Report sections (paste AI-generated content)
        st.subheader("Step 2: Paste Report Content")
        st.info("📝 Copy and paste each section from your AI-generated report below. The formatting will be preserved exactly as you paste it.")
        
        # Report title section
        report_title = st.text_area(
            "Report of Guardian Ad Litem",
            height=100,
            placeholder="Paste the introductory paragraph here...",
            key="report_title_section",
            help="This is the paragraph right below Report of Guardian Ad Litem"
        )
        
        # Sources Consulted section
        sources_consulted = st.text_area(
            "SOURCES CONSULTED:",
            height=150,
            placeholder="Paste the entire SOURCES CONSULTED section here...",
            key="report_sources_section",
            help="This section describes what documents were reviewed"
        )
        
        # Facts section
        facts_section = st.text_area(
            "FACTS:",
            height=200,
            placeholder="Paste the entire FACTS section here...",
            key="report_facts_section",
            help="This section contains the background and payment details"
        )
        
        # Valuation and Recommendation section
        valuation_section = st.text_area(
            "VALUATION and RECOMMENDATION:",
            height=200,
            placeholder="Paste the entire VALUATION and RECOMMENDATION section here...",
            key="report_valuation_section",
            help="This section contains the analysis and final recommendation"
        )
        
        st.write("---")
        
        # STEP 3: Generate formatted report
        st.subheader("Step 3: Generate Formatted Report")
        
        if st.button("Generate Formatted Report", key="format_report_button"):
            # Check that all required fields are filled
            if not all([cause_number, factoring_company, courthouse, report_title, sources_consulted, facts_section, valuation_section]):
                st.error("❌ Please fill in all required fields before generating the report.")
            else:
                st.success("✅ Report formatted successfully!")
                
                # Generate the Word document using the old perfect formatting
                word_doc = create_tombs_maxwell_template(
                    cause_number, factoring_company, courthouse, 
                    report_title, sources_consulted, facts_section, 
                    valuation_section
                )
                
                if word_doc:
                    # Get current year's last two digits
                    current_year = datetime.now().year
                    year_suffix = str(current_year)[-2:]  # Get last 2 digits (e.g., "25" for 2025)
                    
                    # Clean factoring company name for filename (remove special characters)
                    clean_company_name = "".join(c for c in factoring_company if c.isalnum() or c in (' ', '-', '_')).strip()
                    clean_company_name = clean_company_name.replace(' ', '_')
                    
                    st.success("✅ Word document generated using Tombs Maxwell template!")
                    st.download_button(
                        label="📥 Download Word Document",
                        data=word_doc,
                        file_name=f"Ad_Litem_{clean_company_name}_{year_suffix}_XX_Draft_1.0.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Also show the text version for backup
                    st.write("---")
                    st.subheader("📋 Text Version (Backup)")
                    st.write("Copy and paste the text below if needed:")
                    
                    # Generate the complete formatted text report
                    formatted_report = f"""CAUSE NO. {cause_number}

IN RE:

{factoring_company}

{courthouse}

{report_title}

{sources_consulted}

{facts_section}

{valuation_section}

Respectfully submitted,

/s/ Joseph W. Tombs
Joseph W. Tombs
TOMBS MAXWELL, LLP
State Bar No. 20116250
7021 Kewanee Ave. 7-102
Lubbock, TX 79424
Office (806) 698-1122
Tombs@tombsmaxwell.com"""
                    
                    # Display the formatted report in a text area for easy copying
                    st.text_area(
                        "Formatted Report Text", 
                        formatted_report, 
                        height=400, 
                        key="formatted_report_output"
                    )
                    
                else:
                    st.error("❌ Error generating Word document. Please try again.")
                    
                    # Fallback to text-only version
                    st.write("**📋 Text Version Available:**")
                    formatted_report = f"""CAUSE NO. {cause_number}

IN RE:

{factoring_company}

{courthouse}

{report_title}

{sources_consulted}

{facts_section}

{valuation_section}

Respectfully submitted,

/s/ Joseph W. Tombs
Joseph W. Tombs
TOMBS MAXWELL, LLP
State Bar No. 20116250
7021 Kewanee Ave. 7-102
Lubbock, TX 79424
Office (806) 698-1122
Tombs@tombsmaxwell.com"""
                    
                    st.text_area(
                        "Formatted Report Text", 
                        formatted_report, 
                        height=400, 
                        key="formatted_report_fallback"
                    )
